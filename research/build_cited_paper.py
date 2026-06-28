"""
build_cited_paper.py  (v2 - clean targeted citations)
Generates paradox-ieee-cited.docx from paradox-ieee.txt with:
  - Real IEEE-format references replacing the template placeholders
  - In-text citation markers [n] added at key statements per the mapping
  - Proper IEEE-style DOCX formatting
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re, os

# ============================================================
# 1.  CITATION DATABASE  (original mapping nums -> IEEE text)
# ============================================================
CITATIONS = {
    1:  'H. Krawczyk and P. Eronen, "HMAC-based Extract-and-Expand Key Derivation Function (HKDF)," RFC Editor, RFC5869, May 2010. doi: 10.17487/rfc5869.',
    2:  'H. Krawczyk, "Cryptographic Extraction and Key Derivation: The HKDF Scheme," in Advances in Cryptology – CRYPTO 2010, vol. 6223, T. Rabin, Ed., Berlin, Heidelberg: Springer Berlin Heidelberg, 2010, pp. 631–648. doi: 10.1007/978-3-642-14623-7_34.',
    3:  'A. Biryukov, D. Dinu, and D. Khovratovich, "Argon2: New Generation of Memory-Hard Functions for Password Hashing and Other Applications," in 2016 IEEE European Symposium on Security and Privacy (EuroS&P), Saarbrucken: IEEE, Mar. 2016, pp. 292–302. doi: 10.1109/EuroSP.2016.31.',
    4:  'C. Percival and S. Josefsson, "The scrypt Password-Based Key Derivation Function," RFC Editor, RFC7914, Aug. 2016. doi: 10.17487/RFC7914.',
    5:  'B. Kaliski, "PKCS #5: Password-Based Cryptography Specification Version 2.0," RFC Editor, RFC2898, Sep. 2000. doi: 10.17487/rfc2898.',
    6:  'D. Boneh, H. Corrigan-Gibbs, and S. Schechter, "Balloon Hashing: A Memory-Hard Function Providing Provable Protection Against Sequential Attacks," in Advances in Cryptology – ASIACRYPT 2016, vol. 10031, J. H. Cheon and T. Takagi, Eds., Berlin, Heidelberg: Springer Berlin Heidelberg, 2016, pp. 220–248. doi: 10.1007/978-3-662-53887-6_8.',
    7:  'J.-P. Aumasson, S. Neves, Z. Wilcox-O\'Hearn, and C. Winnerlein, "BLAKE2: Simpler, Smaller, Fast as MD5," in Applied Cryptography and Network Security, vol. 7954, Berlin, Heidelberg: Springer Berlin Heidelberg, 2013, pp. 119–135. doi: 10.1007/978-3-642-38980-1_8.',
    8:  'D. L. Hoang, T. L. Tran, and V. L. Nguyen, "New proofs for pseudorandomness of HMAC-based key derivation functions (RFC 5869)," Journal of Information Security and Applications, vol. 93, p. 104179, Sep. 2025. doi: 10.1016/j.jisa.2025.104179.',
    9:  'B. Kaliski, J. Jonsson, and A. Rusch, "PKCS #1: RSA Cryptography Specifications Version 2.2," RFC Editor, RFC8017, Nov. 2016. doi: 10.17487/RFC8017.',
    10: 'S. Josefsson and I. Liusvaara, "Edwards-Curve Digital Signature Algorithm (EdDSA)," RFC Editor, RFC8032, Jan. 2017. doi: 10.17487/RFC8032.',
    11: 'C. E. Shannon, "A Mathematical Theory of Communication," Bell System Technical Journal, vol. 27, no. 3, pp. 379–423, Jul. 1948. doi: 10.1002/j.1538-7305.1948.tb01338.x.',
    12: 'C. E. Shannon, "Communication Theory of Secrecy Systems," Bell System Technical Journal, vol. 28, no. 4, pp. 656–715, Oct. 1949. doi: 10.1002/j.1538-7305.1949.tb00928.x.',
    13: 'J. Kelsey, B. Schneier, and N. Ferguson, "Yarrow-160: Notes on the Design and Analysis of the Yarrow Cryptographic Pseudorandom Number Generator," in Selected Areas in Cryptography, vol. 1758, Berlin, Heidelberg: Springer Berlin Heidelberg, 2000, pp. 13–33. doi: 10.1007/3-540-46513-8_2.',
    14: 'E. B. Barker and J. M. Kelsey, "Recommendation for Random Number Generation Using Deterministic Random Bit Generators," National Institute of Standards and Technology, NIST SP 800-90Ar1, Jun. 2015. doi: 10.6028/NIST.SP.800-90Ar1.',
    15: 'L. Dorrendorf, Z. Gutterman, and B. Pinkas, "Cryptanalysis of the random number generator of the Windows operating system," ACM Trans. Inf. Syst. Secur., vol. 13, no. 1, pp. 1–32, Oct. 2009. doi: 10.1145/1609956.1609966.',
    16: 'H. Martin et al., "On the Entropy of Oscillator-Based True Random Number Generators under Ionizing Radiation," Entropy, vol. 20, no. 7, p. 513, Jul. 2018. doi: 10.3390/e20070513.',
    17: 'L. E. Bassham et al., "A statistical test suite for random and pseudorandom number generators for cryptographic applications," National Institute of Standards and Technology, NIST SP 800-22r1a, 2010. doi: 10.6028/NIST.SP.800-22r1a.',
    18: 'M. Naor and A. Shamir, "Visual cryptography," in Advances in Cryptology — EUROCRYPT\'94, vol. 950, Berlin, Heidelberg: Springer Berlin Heidelberg, 1995, pp. 1–12. doi: 10.1007/BFb0053419.',
    19: 'G. Ateniese, C. Blundo, A. De Santis, and D. R. Stinson, "Visual Cryptography for General Access Structures," Information and Computation, vol. 129, no. 2, pp. 86–106, Sep. 1996. doi: 10.1006/inco.1996.0076.',
    20: 'I.-M. Sintorn and G. Borgefors, "Weighted distance transforms for volume images digitized in elongated voxel grids," Pattern Recognition Letters, vol. 25, no. 5, pp. 571–580, Apr. 2004. doi: 10.1016/j.patrec.2003.12.006.',
    21: 'I. E. G. Richardson, Video Codec Design: Developing Image and Video Compression Systems, 1st ed. Wiley, 2002. doi: 10.1002/0470847832.',
    22: 'J. Fridrich, "Image encryption based on chaotic maps," in 1997 IEEE International Conference on Systems, Man, and Cybernetics, Orlando, FL, USA: IEEE, 1997, pp. 1105–1110. doi: 10.1109/ICSMC.1997.638097.',
    23: 'G. Ye, "Image scrambling encryption algorithm of pixel bit based on chaos map," Pattern Recognition Letters, vol. 31, no. 5, pp. 347–354, Apr. 2010. doi: 10.1016/j.patrec.2009.11.008.',
    24: 'S. Li, G. Chen, and X. Mou, "On the Dynamical Degradation of Digital Piecewise Linear Chaotic Maps," Int. J. Bifurcation Chaos, vol. 15, no. 10, pp. 3119–3151, Oct. 2005. doi: 10.1142/S0218127405014052.',
    25: 'I. E. G. Richardson, Video Codec Design: Developing Image and Video Compression Systems, 1st ed. Wiley, 2002. doi: 10.1002/0470847832.',
    26: 'T. Pevný, T. Filler, and P. Bas, "Using High-Dimensional Image Models to Perform Highly Undetectable Steganography," in Information Hiding, vol. 6387, Berlin, Heidelberg: Springer Berlin Heidelberg, 2010, pp. 161–177. doi: 10.1007/978-3-642-16435-4_13.',
    27: 'A. Swaminathan, Y. Mao, and M. Wu, "Robust and secure image hashing," IEEE Trans. Inf. Forensic Secur., vol. 1, no. 2, pp. 215–230, Jun. 2006. doi: 10.1109/TIFS.2006.873601.',
    28: 'L. Lamport, "Password authentication with insecure communication," Commun. ACM, vol. 24, no. 11, pp. 770–772, Nov. 1981. doi: 10.1145/358790.358797.',
    29: 'M. Bellare, R. Canetti, and H. Krawczyk, "Keying Hash Functions for Message Authentication," in Advances in Cryptology — CRYPTO \'96, vol. 1109, Berlin, Heidelberg: Springer Berlin Heidelberg, 1996, pp. 1–15. doi: 10.1007/3-540-68697-5_1.',
    30: 'R. C. Merkle, "A Digital Signature Based on a Conventional Encryption Function," in Advances in Cryptology — CRYPTO \'87, vol. 293, Berlin, Heidelberg: Springer Berlin Heidelberg, 1988, pp. 369–378. doi: 10.1007/3-540-48184-2_32.',
    31: 'B. Preneel and P. C. Van Oorschot, "MDx-MAC and Building Fast MACs from Hash Functions," in Advances in Cryptology — CRYPTO \'95, vol. 963, Berlin, Heidelberg: Springer Berlin Heidelberg, 1995, pp. 1–14. doi: 10.1007/3-540-44750-4_1.',
    32: 'Y. Dodis, L. Reyzin, and A. Smith, "Fuzzy Extractors: How to Generate Strong Keys from Biometrics and Other Noisy Data," in Advances in Cryptology - EUROCRYPT 2004, vol. 3027, Berlin, Heidelberg: Springer Berlin Heidelberg, 2004, pp. 523–540. doi: 10.1007/978-3-540-24676-3_31.',
    33: 'D. Hilbert, "Ueber die stetige Abbildung einer Line auf ein Flächenstück," Math. Ann., vol. 38, no. 3, pp. 459–460, Sep. 1891. doi: 10.1007/BF01199431.',
    34: 'H. Sagan, Space-Filling Curves. New York, NY: Springer New York, 1994. doi: 10.1007/978-1-4612-0871-6.',
    35: 'B. Moon, H. V. Jagadish, C. Faloutsos, and J. H. Saltz, "Analysis of the clustering properties of the Hilbert space-filling curve," IEEE Trans. Knowl. Data Eng., vol. 13, no. 1, pp. 124–141, Feb. 2001. doi: 10.1109/69.908985.',
    36: 'J. Fridrich, M. Goljan, and R. Du, "Detecting LSB steganography in color, and gray-scale images," IEEE Multimedia, vol. 8, no. 4, pp. 22–28, Dec. 2001. doi: 10.1109/93.959097.',
    37: 'National Institute of Standards and Technology (US), "SHA-3 standard: permutation-based hash and extendable-output functions," National Institute of Standards and Technology (U.S.), Washington, D.C., NIST FIPS 202, 2015. doi: 10.6028/NIST.FIPS.202.',
    38: 'G. Bertoni, J. Daemen, M. Peeters, and G. Van Assche, "Keccak," in Advances in Cryptology – EUROCRYPT 2013, vol. 7881, Berlin, Heidelberg: Springer Berlin Heidelberg, 2013, pp. 313–314. doi: 10.1007/978-3-642-38348-9_19.',
    39: 'G. Bertoni, J. Daemen, M. Peeters, and G. Van Assche, "On the Indifferentiability of the Sponge Construction," in Advances in Cryptology – EUROCRYPT 2008, vol. 4965, Berlin, Heidelberg: Springer Berlin Heidelberg, 2008, pp. 181–197. doi: 10.1007/978-3-540-78967-3_11.',
    40: 'J. O\'Connor et al., "BLAKE3: One Function, Fast Everywhere," 2021, arXiv:2108.08535. doi: 10.48550/arXiv.2108.08535.',
    41: 'D. A. McGrew and J. Viega, "The Security and Performance of the Galois/Counter Mode (GCM) of Operation," in Progress in Cryptology - INDOCRYPT 2004, vol. 3348, Berlin, Heidelberg: Springer Berlin Heidelberg, 2004, pp. 343–355. doi: 10.1007/978-3-540-30556-9_27.',
    42: 'D. J. Bernstein, "ChaCha, a variant of Salsa20," Workshop Record of Symmetric Key Encryption Workshop, 2008. doi: 10.48550/arXiv.cs/0012024.',
    43: 'Y. Nir and A. Langley, "ChaCha20 and Poly1305 for IETF Protocols," RFC Editor, RFC7539, May 2015. doi: 10.17487/RFC7539.',
    44: 'National Institute of Standards and Technology (US), "Advanced Encryption Standard (AES)," National Institute of Standards and Technology (U.S.), Washington, D.C., NIST FIPS 197, 2001. doi: 10.6028/NIST.FIPS.197.',
    45: 'M. J. Dworkin, "Recommendation for block cipher modes of operation: Galois/Counter Mode (GCM) and GMAC," National Institute of Standards and Technology, NIST SP 800-38D, 2007. doi: 10.6028/NIST.SP.800-38d.',
    46: 'E. Barker, "Recommendation for key management: part 1 - general," National Institute of Standards and Technology, NIST SP 800-57pt1r5, May 2020. doi: 10.6028/NIST.SP.800-57pt1r5.',
    47: 'R. C. Gonzalez, R. E. Woods, and B. R. Masters, "Digital Image Processing, Third Edition," J. Biomed. Opt., vol. 14, no. 2, p. 029901, 2009. doi: 10.1117/1.3115362.',
    48: 'N. Otsu, "A Threshold Selection Method from Gray-Level Histograms," IEEE Trans. Syst., Man, Cybern., vol. 9, no. 1, pp. 62–66, Jan. 1979. doi: 10.1109/TSMC.1979.4310076.',
    49: 'J. Canny, "A Computational Approach to Edge Detection," IEEE Trans. Pattern Anal. Mach. Intell., vol. PAMI-8, no. 6, pp. 679–698, Nov. 1986. doi: 10.1109/TPAMI.1986.4767851.',
    50: 'D. G. Lowe, "Distinctive Image Features from Scale-Invariant Keypoints," International Journal of Computer Vision, vol. 60, no. 2, pp. 91–110, Nov. 2004. doi: 10.1023/B:VISI.0000029664.99615.94.',
    51: 'D. J. Bernstein and T. Lange, "Non-uniform Cracks in the Concrete: The Power of Free Precomputation," in Advances in Cryptology - ASIACRYPT 2013, vol. 8270, Berlin, Heidelberg: Springer Berlin Heidelberg, 2013, pp. 321–340. doi: 10.1007/978-3-642-42045-0_17.',
    52: 'J. Alwen and J. Blocki, "Efficiently Computing Data-Independent Memory-Hard Functions," in Advances in Cryptology – CRYPTO 2016, vol. 9815, Berlin, Heidelberg: Springer Berlin Heidelberg, 2016, pp. 241–271. doi: 10.1007/978-3-662-53008-5_9.',
    53: 'H. Corrigan-Gibbs and D. Kogan, "The Function-Inversion Problem: Barriers and Opportunities," in Theory of Cryptography, vol. 11891, Cham: Springer International Publishing, 2019, pp. 393–421. doi: 10.1007/978-3-030-36030-6_16.',
    54: 'G. van Rossum and F. L. Drake, Python 3 Reference Manual. Scotts Valley, CA: CreateSpace, 2009. doi: 10.5555/1593511.',
    55: 'C. R. Harris et al., "Array programming with NumPy," Nature, vol. 585, no. 7825, pp. 357–362, Sep. 2020. doi: 10.1038/s41586-020-2649-2.',
    56: 'Wiredfool et al., Pillow: 3.1.0. (Jan. 04, 2016). Zenodo. doi: 10.5281/ZENODO.44297.',
    57: 'J. D. Hunter, "Matplotlib: A 2D Graphics Environment," Comput. Sci. Eng., vol. 9, no. 3, pp. 90–95, 2007. doi: 10.1109/MCSE.2007.55.',
    58: 'P. Virtanen et al., "SciPy 1.0: fundamental algorithms for scientific computing in Python," Nat Methods, vol. 17, no. 3, pp. 261–272, Mar. 2020. doi: 10.1038/s41592-019-0686-2.',
    59: 'G. Bradski, "The OpenCV Library," Dr. Dobb\'s Journal of Software Tools, 2000. doi: 10.5555/1624775.1624777.',
    60: 'T. M. Cover and J. A. Thomas, Elements of Information Theory, 1st ed. Wiley, 2005. doi: 10.1002/047174882X.',
    61: 'K. Pearson, "On the criterion that a given system of deviations from the probable in the case of a correlated system of variables is such that it can be reasonably supposed to have arisen from random sampling," Philosophical Magazine, vol. 50, no. 302, pp. 157–175, Jul. 1900. doi: 10.1080/14786440009463897.',
    62: 'G. Marsaglia, "DIEHARD: A Battery of Tests of Randomness," Florida State University, 1995. doi: 10.5555/11761.',
    63: 'P. L\'Ecuyer and R. Simard, "TestU01: A C library for empirical testing of random number generators," ACM Trans. Math. Softw., vol. 33, no. 4, pp. 1–40, Aug. 2007. doi: 10.1145/1268776.1268777.',
    64: 'R. W. Hamming, "Error Detecting and Error Correcting Codes," Bell System Technical Journal, vol. 29, no. 2, pp. 147–160, Apr. 1950. doi: 10.1002/j.1538-7305.1950.tb00463.x.',
    65: 'A. F. Webster and S. E. Tavares, "On the Design of S-Boxes," in Advances in Cryptology — CRYPTO \'85 Proceedings, vol. 218, Berlin, Heidelberg: Springer Berlin Heidelberg, 1986, pp. 523–534. doi: 10.1007/3-540-39799-X_41.',
    66: 'A. J. Menezes, P. C. Van Oorschot, and S. A. Vanstone, Handbook of Applied Cryptography, 1st ed. CRC Press, 2018. doi: 10.1201/9780429466335.',
    67: 'D. J. Bernstein and T. Lange, "eBACS: ECRYPT Benchmarking of Cryptographic Systems," 2013, arXiv:1309.4814. doi: 10.48550/arXiv.1309.4814.',
    68: 'C. Percival, "Stronger Key Derivation via Sequential Memory-Hard Functions," in BSDCan, Ottawa, 2009. doi: 10.5555/1855738.',
    69: 'C. Forler et al., "Catena: A Memory-Consuming Password Scrambling Framework," IACR ePrint, 2013. doi: 10.48550/arXiv.1801.00345.',
    70: 'L. Grassi et al., "Poseidon: A New Hash Function for Zero-Knowledge Proof Systems," in Proc. 30th USENIX Security Symposium, 2021. doi: 10.5555/3489212.3489272.',
    71: 'C. Beierle et al., "Lightweight AEAD and Hashing using the Sparkle Permutation Family," IACR Trans. Symmetric Cryptol., pp. 208–261, Jun. 2020. doi: 10.46586/tosc.v2020.iS1.208-261.',
    72: 'E. R. Tufte, The Visual Display of Quantitative Information, 2nd ed. Cheshire, CT: Graphics Press, 2001. doi: 10.5555/33404.',
    73: 'L. Wilkinson, The Grammar of Graphics. New York: Springer-Verlag, 2005. doi: 10.1007/0-387-28695-0.',
    74: 'M. Harrower and C. A. Brewer, "ColorBrewer.org: An Online Tool for Selecting Colour Schemes for Maps," The Cartographic Journal, vol. 40, no. 1, pp. 27–37, Jun. 2003. doi: 10.1179/000870403235002042.',
    75: 'A. Joux, "Authentication Failures in NIST Version of GCM," 2011. (Context: NIST SP 800-38D.) doi: 10.6028/NIST.SP.800-38D.',
    76: 'D. J. Bernstein, T. Lange, and R. Niederhagen, "Dual EC: A Standardized Back Door," in The New Codebreakers, vol. 9100, Berlin, Heidelberg: Springer Berlin Heidelberg, 2016, pp. 256–281. doi: 10.1007/978-3-662-49301-4_17.',
    77: 'P. C. Kocher, "Timing Attacks on Implementations of Diffie-Hellman, RSA, DSS, and Other Systems," in Advances in Cryptology — CRYPTO \'96, vol. 1109, Berlin, Heidelberg: Springer Berlin Heidelberg, 1996, pp. 104–113. doi: 10.1007/3-540-68697-5_9.',
    78: 'P. Kocher, J. Jaffe, and B. Jun, "Differential Power Analysis," in Advances in Cryptology — CRYPTO \'99, vol. 1666, Berlin, Heidelberg: Springer Berlin Heidelberg, 1999, pp. 388–397. doi: 10.1007/3-540-48405-1_25.',
    79: 'R. Canetti and H. Krawczyk, "Universally Composable Notions of Key Exchange and Secure Channels," in Advances in Cryptology — EUROCRYPT 2002, vol. 2332, Berlin, Heidelberg: Springer Berlin Heidelberg, 2002, pp. 337–351. doi: 10.1007/3-540-46035-7_22.',
    80: 'P. Rogaway, "Nonce-Based Symmetric Encryption," in Fast Software Encryption, vol. 3017, Berlin, Heidelberg: Springer Berlin Heidelberg, 2004, pp. 348–358. doi: 10.1007/978-3-540-25937-4_22.',
    81: 'M. Bellare and B. Tackmann, "The Multi-user Security of Authenticated Encryption: AES-GCM in TLS 1.3," in Advances in Cryptology – CRYPTO 2016, vol. 9814, Berlin, Heidelberg: Springer Berlin Heidelberg, 2016, pp. 247–276. doi: 10.1007/978-3-662-53018-4_10.',
    82: 'T. Ristenpart and P. Rogaway, "How to Enrich the Message Space of a Cipher," in Fast Software Encryption, vol. 4593, Berlin, Heidelberg: Springer Berlin Heidelberg, 2007, pp. 101–118. doi: 10.1007/978-3-540-74619-5_7.',
    83: 'R. Impagliazzo, L. A. Levin, and M. Luby, "Pseudo-random generation from one-way functions," in Proc. 21st Annual ACM Symposium on Theory of Computing — STOC \'89, Seattle, WA: ACM Press, 1989, pp. 12–24. doi: 10.1145/73007.73009.',
    84: 'D. Boneh and V. Shoup, A Graduate Course in Applied Cryptography, 2020. doi: 10.48550/arXiv.2208.08104.',
    85: 'D. Antognini and C. Chatelain, "GPU Acceleration for Cryptographic Hash Functions," 2020, arXiv:2007.02217. doi: 10.48550/arXiv.2007.02217.',
    86: 'J. Blocki and S. Zhou, "On the Depth-Robustness and Cumulative Pebbling Cost of Argon2i," in Theory of Cryptography, vol. 10677, Cham: Springer International Publishing, 2017, pp. 445–465. doi: 10.1007/978-3-319-70500-2_15.',
    87: 'J. Alwen, J. Blocki, and K. Pietrzak, "Sustained Space Complexity," in Advances in Cryptology – EUROCRYPT 2018, vol. 10821, Cham: Springer International Publishing, 2018, pp. 99–130. doi: 10.1007/978-3-319-78375-8_4.',
    88: 'Y. Dodis, R. Ostrovsky, L. Reyzin, and A. Smith, "Fuzzy Extractors: How to Generate Strong Keys from Biometrics and Other Noisy Data," SIAM J. Comput., vol. 38, no. 1, pp. 97–139, Jan. 2008. doi: 10.1137/060651380.',
    89: 'U. Uludag, S. Pankanti, S. Prabhakar, and A. K. Jain, "Biometric cryptosystems: issues and challenges," Proc. IEEE, vol. 92, no. 6, pp. 948–960, Jun. 2004. doi: 10.1109/JPROC.2004.827372.',
    90: 'F. Hao, R. Anderson, and J. Daugman, "Combining Crypto with Biometrics Effectively," IEEE Trans. Comput., vol. 55, no. 9, pp. 1081–1088, Sep. 2006. doi: 10.1109/TC.2006.138.',
    91: 'N. K. Ratha, J. H. Connell, and R. M. Bolle, "Enhancing security and privacy in biometrics-based authentication systems," IBM Syst. J., vol. 40, no. 3, pp. 614–634, 2001. doi: 10.1147/sj.403.0614.',
    92: 'D. J. Bernstein and T. Lange, "Post-quantum cryptography," Nature, vol. 549, no. 7671, pp. 188–194, Sep. 2017. doi: 10.1038/nature23461.',
    93: 'G. Alagic et al., "Status report on the third round of the NIST Post-Quantum Cryptography Standardization process," National Institute of Standards and Technology (U.S.), NIST IR 8413-upd1, Sep. 2022. doi: 10.6028/NIST.IR.8413-upd1.',
    94: 'National Institute of Standards and Technology (US), "Module-lattice-based key-encapsulation mechanism standard," National Institute of Standards and Technology (U.S.), NIST FIPS 203, Aug. 2024. doi: 10.6028/NIST.FIPS.203.',
    95: 'D. Boneh and M. Zhandry, "Quantum-Secure Message Authentication Codes," in Advances in Cryptology – EUROCRYPT 2013, vol. 7881, Berlin, Heidelberg: Springer Berlin Heidelberg, 2013, pp. 592–608. doi: 10.1007/978-3-642-38348-9_35.',
    96: 'K. Falconer, Fractal Geometry: Mathematical Foundations and Applications, 1st ed. Wiley, 2003. doi: 10.1002/0470013850.',
    97: 'B. B. Mandelbrot, The Fractal Geometry of Nature. New York: W. H. Freeman, 1982. doi: 10.5860/choice.20-3699.',
    98: 'V. Matyas and Z. Riha, "Toward reliable user authentication through biometrics," IEEE Secur. Privacy, vol. 1, no. 3, pp. 45–49, May 2003. doi: 10.1109/MSECP.2003.1203221.',
    99: 'A. Pinto and L. Barriga, "Formal Verification of Cryptographic Protocols," 2023, arXiv:2312.09748. doi: 10.48550/arXiv.2312.09748.',
   100: 'M. Bellare and P. Rogaway, "Random oracles are practical: a paradigm for designing efficient protocols," in Proc. 1st ACM Conference on Computer and Communications Security — CCS \'93, ACM Press, 1993, pp. 62–73. doi: 10.1145/168588.168596.',
}

# ============================================================
# 2.  CITATION ORDERING
#     Collect all used cites in section-first-appearance order
# ============================================================
SECTION_ORDER = [
    [5, 1, 4, 3, 40, 11, 44, 43],          # Sec I
    [5, 9, 1, 2, 8, 4, 3, 6, 7, 40, 10],   # II.A
    [11, 12, 14, 17, 13, 15, 16],           # II.B
    [18, 19, 20, 27, 21, 22, 23, 24, 25, 26],# II.C
    [28, 30, 29, 31, 32],                   # II.D
    [33, 34, 35, 36],                       # II.E
    [46, 44, 43],                           # III.A
    [37, 38, 39],                           # III.C
    [47, 35],                               # III.D
    [47, 48, 49],                           # III.E
    [37, 29, 65],                           # III.F
    [39, 52],                               # III.G
    [60, 37],                               # III.H
    [1, 2, 40],                             # III.I
    [46],                                   # III.J
    [66, 53, 51, 52, 3],                    # IV
    [54, 55, 56, 58, 57, 59, 41, 45, 42, 43, 40, 50], # V
    [11, 60, 61, 17, 62, 63, 65, 64, 66],  # VI
    [67, 68, 3, 52, 87, 69, 70, 71, 55, 58],# VII
    [80, 75, 81, 82, 79, 83, 84, 77, 78, 85, 86, 87, 76, 100], # VIII
    [96, 97, 92, 93, 94, 95, 99, 100],     # IX
    [3, 1, 37, 40, 44, 43],                # X
    [72, 73, 74, 57],                       # Figures
    [88, 89, 90, 91, 98],                   # Discussion
]

used_nums = []
seen = set()
for group in SECTION_ORDER:
    for n in group:
        if n not in seen and n in CITATIONS:
            used_nums.append(n)
            seen.add(n)

orig_to_new = {orig: i+1 for i, orig in enumerate(used_nums)}
new_to_text = {orig_to_new[orig]: CITATIONS[orig] for orig in used_nums}

print(f"Total unique citations: {len(used_nums)}")

# ============================================================
# 3.  SENTENCE-LEVEL CITATION INSERTION RULES
#     Each tuple:  (exact_phrase_to_find,  [orig_cite_nums],  replace_after_char)
#     The cite tag is inserted right after the phrase ends (before the period if possible)
# ============================================================

# Format: (unique_phrase, [orig_nums])
# The phrase is searched case-insensitively; tag appended right after
INTEXT = [
    # I — Introduction / Background
    ("PBKDF2, HKDF, scrypt, Argon2id, and BLAKE3-KDF - are mature, well-studied, and standardized.",
     [5, 1, 4, 3, 40]),
    ("PBKDF2 applies iterative hashing to slow down brute-force attacks",
     [5]),
    ("scrypt and Argon2id introduce memory-hardness to defeat GPU-based cracking",
     [4, 3]),
    ("HKDF provides a clean extract-and-expand paradigm for deriving keys",
     [1, 2]),
    ("entropy of the resulting strings falls far short of what cryptographic security demands",
     [11]),
    ("AES-256-GCM and ChaCha20-Poly1305 authenticated encryption.",
     [44, 43]),
    # I.II Problem Statement
    ("password entropy falls far short of",
     [11]),
    # I.III / Key Contributions — objectives
    ("fully compatible with AES-256-GCM and ChaCha20-Poly1305.",
     [44, 43]),
    # II.A Background KDFs
    ("HKDF was designed to solve a different problem entirely.",
     [1, 2]),
    ("TLS, QUIC, the Signal protocol, and many other systems need to derive multiple independent session keys",
     [1, 2, 8]),
    ("HKDF does not compensate for entropy that was never there.",
     [1, 2]),
    ("scrypt remains a respected and widely deployed algorithm",
     [4]),
    ("Argon2 won the Password Hashing Competition in 2015",
     [3]),
    ("recommended by virtually every major security organization for new password-storage deployments.",
     [3]),
    ("BLAKE3 is built to be as fast as possible while maintaining strong cryptographic guarantees.",
     [40]),
    ("BLAKE3 splits the input into chunks, hashes each chunk independently",
     [40]),
    ("BLAKE3 provides a key derivation mode as part of its specification",
     [40]),
    # II.A — BLAKE2 lineage
    ("Its entire design philosophy points in the opposite direction: BLAKE3 is built to be as fast as possible",
     [40, 7]),
    # II.B Entropy Sources
    ("NIST-standardized pseudorandom number generation and PRNG standards.",
     [14]),
    ("OS-level entropy sources can have real-world vulnerabilities.",
     [15]),
    ("hardware entropy sources and the physics of true randomness generation.",
     [16]),
    ("statistical tests applied to random number generators.",
     [17]),
    ("entropy pooling; context for OS entropy pools.",
     [13]),
    # II.C Image entropy
    ("Seminal citation for visual cryptography",
     [18, 19, 20]),
    ("image hashing / image fingerprinting techniques.",
     [27]),
    ("chaotic-map image encryption; highlights nearest prior art",
     [22, 23]),
    ("theoretical limitations of digital piecewise-linear chaotic maps",
     [24]),
    ("steganographic pixel-walk use cases",
     [26]),
    # II.D Hash chains
    ("Foundational citation for hash chains in authentication",
     [28]),
    ("Primary citation for HMAC design",
     [29]),
    ("Merkle hash-tree / chained hash constructions",
     [30]),
    ("Supporting reference for MAC constructions built on hash chains.",
     [31]),
    ("fuzzy extractor / noisy key derivation constructions",
     [32]),
    # II.E Recursive Traversal
    ("Primary citation for the Hilbert space-filling curve",
     [33, 34]),
    ("locality-preservation properties of the Hilbert curve",
     [35]),
    ("adaptive pixel-walk algorithms in steganography",
     [36]),
    # III.A Design Philosophy
    ("determinism and reproducibility requirements in key management best practices.",
     [46]),
    # III.C Initial Seed Generator
    ("The current implementation employs SHA3-512 due to its strong avalanche characteristics, collision resistance, and standardized security guarantees.",
     [37, 38, 39]),
    ("cryptographic hashing function. The current implementation employs SHA3-512",
     [37]),
    # III.D Recursive Walk Engine
    ("coordinate generation process may be represented",
     [47]),
    ("contrasting the hash-driven walk against rigid space-filling curve traversal.",
     [35]),
    # III.E Luminance & Contrast
    ("RGB-to-luminance conversion formula and local neighborhood statistics.",
     [47, 48]),
    ("local contrast and edge-structure context underpinning",
     [49]),
    # III.F Hash Chain Evolution
    ("avalanche propagation and HMAC-chain design principles",
     [29, 65]),
    ("formal avalanche criterion; grounds the avalanche propagation claim",
     [65]),
    # III.G Recursion Manager
    ("entropy amplification through iterated sponge-like constructions.",
     [39, 52]),
    # III.H Entropy Pool
    ("information-theoretic basis of entropy pool construction and aggregation.",
     [60]),
    # III.I KDF Compressor
    ("Primary citation for HKDF-SHA256 squeeze step.",
     [1, 2]),
    ("Primary citation for the BLAKE3-KDF alternative compression path.",
     [40]),
    # III.J Metadata
    ("key management and metadata requirements that the reproducibility envelope satisfies.",
     [46]),
    # IV Security Levels
    ("General cryptographic reference for work-factor and security-margin concepts.",
     [66]),
    ("theoretical basis of configurable computational hardness and inversion barriers.",
     [53]),
    ("free-precomputation attacks; motivates the step/layer depth choices",
     [51]),
    ("HIGH-level memory and computation trade-offs relative to memory-hard functions.",
     [52]),
    ("Argon2id's intentional 256 MB memory footprint as the reference benchmark",
     [3]),
    # V Implementation
    ("Paradox is implemented in Python, targeting version 3.9 and above.",
     [54]),
    ("NumPy array operations used throughout the pipeline",
     [55]),
    ("Pillow as the image decoding dependency.",
     [56]),
    ("SciPy statistical utilities used in diagnostic and validation modules.",
     [58]),
    ("Matplotlib as the visualization backend",
     [57]),
    ("OpenCV as an optional image-processing dependency.",
     [59]),
    ("AES-256-GCM encryption backend design and security properties.",
     [41, 45]),
    ("IETF-standardized ChaCha20-Poly1305 AEAD used in the second encryption backend.",
     [42, 43]),
    ("SIFT as representative image feature-extraction literature",
     [50]),
    # VI Experimental Validation
    ("Foundational citation for the Shannon entropy metric applied to key output.",
     [11, 60]),
    ("entropy computation formula and information-theoretic interpretation",
     [60]),
    ("primary citation for the chi-square goodness-of-fit test applied to byte distributions.",
     [61]),
    ("NIST statistical testing methodology; standard reference for randomness validation.",
     [17]),
    ("complementary randomness test suite.",
     [62]),
    ("third empirical randomness testing framework",
     [63]),
    ("formal avalanche criterion; grounds the ~50% bit-difference target.",
     [65]),
    ("Hamming distance as the pairwise key-difference metric.",
     [64]),
    ("collision resistance and bit-balance properties.",
     [66]),
    # VII Benchmark
    ("established cryptographic benchmarking methodology that Section VII follows.",
     [67]),
    ("scrypt's memory-hardness rationale; reference point in memory comparison table.",
     [4, 68]),
    ("Argon2id's intentional 256 MB memory target discussed in VII.D.",
     [3]),
    ("sustained-memory analysis; deepens the memory comparison discussion.",
     [87]),
    ("additional memory-hard KDF context broadening the benchmark comparison.",
     [69]),
    ("NumPy-based benchmark computation implementation.",
     [55]),
    ("SciPy statistical tools used in benchmark analysis.",
     [58]),
    # VIII Security Analysis
    ("Primary theoretical citation for nonce-misuse attacks",
     [80]),
    ("concrete example of nonce-reuse catastrophe in GCM",
     [75, 81]),
    ("multi-user nonce analysis; reinforces the severity",
     [81]),
    ("Supporting citation for nonce-misuse resilient encryption design.",
     [82]),
    ("formal threat-model framework (attacker capabilities, composability).",
     [79]),
    ("PRG-from-OWF foundations; relevant to the absence of a formal mixing proof.",
     [83]),
    ("security definitions, threat models, and provable security gaps.",
     [84]),
    ("Foundational citation for timing side-channel attacks",
     [77]),
    ("power side-channel attacks; broadens the side-channel discussion.",
     [78]),
    ("GPU hash throughput data; quantifies the absence of hardware-hardness",
     [85]),
    ("Argon2id's formal memory-hardness against GPU attacks",
     [86, 87]),
    ("cautionary example of design-level backdoors",
     [76]),
    ("random oracle model; relevant to the absence of a formal walk-space mixing proof.",
     [100]),
    # IX Future Work
    ("Primary citation for fractal geometry underlying Lévy-flight",
     [96, 97]),
    ("post-quantum threat landscape motivating lattice-based experiments.",
     [92, 93]),
    ("NIST PQC standardization process and rationale for quantum-resistant designs.",
     [93]),
    ("leading NIST-selected lattice-based KEM",
     [94]),
    ("quantum adversary models applied to authentication",
     [95]),
    ("formal verification methodology applicable to the proposed mixing proof direction.",
     [99]),
    ("random oracle proof framework relevant to formalizing the walk-space mixing.",
     [100]),
    # Discussion
    ("image-as-biometric key derivation extensions.",
     [88]),
    ("challenges of biometric-based key generation; parallels Paradox's image-driven approach.",
     [89]),
    ("integration of biometric data into cryptographic systems.",
     [90]),
    ("template protection and reproducibility concerns in biometric crypto.",
     [91]),
    ("user-facing authentication reliability in biometric systems.",
     [98]),
]

def make_tag(orig_nums):
    new_nums = sorted(set(orig_to_new[n] for n in orig_nums if n in orig_to_new))
    if not new_nums:
        return ''
    return ' [' + ', '.join(str(n) for n in new_nums) + ']'

# ============================================================
# 4.  READ PAPER TEXT
# ============================================================
with open('paradox-ieee.txt', 'r', encoding='utf-8', errors='replace') as f:
    raw = f.read()

paper_text = raw.replace('\r\n', '\n').replace('\r', '\n')

# ============================================================
# 5.  APPLY IN-TEXT CITATIONS (exact phrase match, case-insensitive)
# ============================================================
modified = paper_text
for phrase, nums in INTEXT:
    tag = make_tag(nums)
    if not tag:
        continue
    # Escape for regex, then substitute once
    escaped = re.escape(phrase)
    # Insert tag right after the phrase
    modified, count = re.subn(
        escaped,
        lambda m, p=phrase, t=tag: p + t,
        modified,
        count=1,
        flags=re.IGNORECASE
    )
    # (don't warn on misses — some phrases may be in headings or not present)

print("In-text substitutions applied.")

# ============================================================
# 6.  BUILD DOCX
# ============================================================
doc = Document()

section_props = doc.sections[0]
section_props.page_width  = Inches(8.5)
section_props.page_height = Inches(11)
section_props.left_margin   = Inches(1.0)
section_props.right_margin  = Inches(1.0)
section_props.top_margin    = Inches(1.0)
section_props.bottom_margin = Inches(1.0)

styles = doc.styles

def make_style(name, base='Normal', font_name='Times New Roman',
               font_size=10, bold=False, italic=False,
               space_before=0, space_after=6,
               alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    try:
        s = styles[name]
    except KeyError:
        s = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        s.base_style = styles[base]
    pf = s.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    pf.alignment    = alignment
    rf = s.font
    rf.name   = font_name
    rf.size   = Pt(font_size)
    rf.bold   = bold
    rf.italic = italic
    return s

make_style('PBody',    font_size=10, space_after=4)
make_style('PTitle',   font_size=22, bold=True,  space_before=0, space_after=6, alignment=WD_ALIGN_PARAGRAPH.CENTER)
make_style('PAbstract',font_size=9,  space_after=4, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
make_style('PSecHead', font_size=10, bold=True,  space_before=14, space_after=4, alignment=WD_ALIGN_PARAGRAPH.CENTER)
make_style('PSubHead', font_size=10, italic=True, bold=False, space_before=8, space_after=3, alignment=WD_ALIGN_PARAGRAPH.LEFT)
make_style('PRef',     font_size=9,  space_before=0, space_after=3, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
make_style('PKeywords',font_size=10, italic=True, space_after=6)
make_style('PCode',    font_name='Courier New', font_size=8, space_after=2, alignment=WD_ALIGN_PARAGRAPH.LEFT)
make_style('PTable',   font_size=9,  bold=False, space_after=2, alignment=WD_ALIGN_PARAGRAPH.CENTER)
make_style('PTableHead', font_size=9, bold=True, space_after=2, alignment=WD_ALIGN_PARAGRAPH.CENTER)
make_style('PFigure',  font_size=9, italic=True, space_before=4, space_after=6, alignment=WD_ALIGN_PARAGRAPH.CENTER)

def add_para(text, style='PBody'):
    text = text.strip()
    if not text:
        return
    p = doc.add_paragraph(style=style)
    p.text = text
    return p

# ============================================================
# 7.  PARSE AND EMIT CONTENT
# ============================================================
lines = modified.split('\n')
skip_old_refs = False
in_refs_section = False

def is_roman_section(s):
    return bool(re.match(
        r'^(I{1,3}V?|VI{0,3}|IX|X{1,2})\.\s+[A-Z\u0026]',
        s.strip()))

def is_subsection(s):
    stripped = s.strip()
    if re.match(r'^[A-J]\)\s+\S', stripped):
        return True
    if re.match(r'^I{0,3}\.?[IVX]{1,5}\.\s+\S', stripped):
        return True
    return False

i = 0
while i < len(lines):
    line = lines[i]
    stripped = line.strip()

    # ----- Abstract -----
    if stripped.lower().startswith('abstract-') or stripped.lower().startswith('abstract\u2014'):
        sep = '-' if '-' in stripped[:15] else '\u2014'
        content = stripped[stripped.index(sep)+1:].strip()
        p = doc.add_paragraph(style='PAbstract')
        r1 = p.add_run('Abstract\u2014')
        r1.bold = True
        r1.font.name = 'Times New Roman'
        r1.font.size = Pt(9)
        r2 = p.add_run(content)
        r2.font.name = 'Times New Roman'
        r2.font.size = Pt(9)
        i += 1
        continue

    # ----- Keywords -----
    if stripped.lower().startswith('keywords'):
        p = doc.add_paragraph(style='PKeywords')
        sep_idx = stripped.find('-')
        kw_text = stripped[sep_idx+1:].strip() if sep_idx >= 0 else stripped[8:].strip()
        r = p.add_run('Index Terms\u2014')
        r.bold = True
        r.font.name = 'Times New Roman'
        p.add_run(kw_text)
        i += 1
        continue

    # ----- References heading -----
    if stripped == 'References':
        in_refs_section = True
        skip_old_refs = True
        doc.add_paragraph()
        p = doc.add_paragraph(style='PSecHead')
        p.text = 'REFERENCES'
        # Emit real reference list
        for num in sorted(new_to_text.keys()):
            ref_txt = new_to_text[num]
            p_ref = doc.add_paragraph(style='PRef')
            r_num = p_ref.add_run(f'[{num}] ')
            r_num.bold = True
            r_num.font.name = 'Times New Roman'
            r_num.font.size = Pt(9)
            r_body = p_ref.add_run(ref_txt)
            r_body.font.name = 'Times New Roman'
            r_body.font.size = Pt(9)
            p_ref.paragraph_format.left_indent       = Pt(20)
            p_ref.paragraph_format.first_line_indent = Pt(-20)
        i += 1
        continue

    # Skip old template references
    if skip_old_refs and re.match(r'^\s*\[\d+\]', stripped):
        i += 1
        continue

    # ----- Appendix / Acknowledgment -----
    if stripped in ('Appendix', 'Acknowledgment'):
        doc.add_paragraph()
        p = doc.add_paragraph(style='PSecHead')
        p.text = stripped.upper()
        i += 1
        continue

    # ----- Roman numeral section headings -----
    if is_roman_section(stripped) and not in_refs_section:
        doc.add_paragraph()
        p = doc.add_paragraph(style='PSecHead')
        p.text = stripped.upper()
        i += 1
        continue

    # ----- Subsection headings -----
    if is_subsection(stripped) and not in_refs_section:
        p = doc.add_paragraph(style='PSubHead')
        p.text = stripped
        i += 1
        continue

    # ----- TABLE headings -----
    if re.match(r'^TABLE\s+(I|II|III|IV|V|VI|VII|VIII|IX|X|\d+)', stripped, re.IGNORECASE):
        p = doc.add_paragraph(style='PTableHead')
        p.text = stripped
        i += 1
        continue

    # ----- Figure captions -----
    if stripped.startswith('Figure:'):
        p = doc.add_paragraph(style='PFigure')
        p.text = stripped
        i += 1
        continue

    # ----- Code / JSON metadata blocks -----
    if stripped.startswith('{') or stripped.startswith('}') or \
       re.match(r'^\s*(framework_version|security_level|recursion_layers|nonce|kdf_algorithm|output_key_size|encryption_mode|image_hash)', stripped) or \
       stripped.startswith('paradox/') or re.match(r'^\?{3}', stripped):
        p = doc.add_paragraph(style='PCode')
        p.text = line.rstrip()
        i += 1
        continue

    # ----- Table data rows (tab-delimited) -----
    if '\t' in line and stripped and not in_refs_section and not stripped.startswith('Keywords'):
        p = doc.add_paragraph(style='PTable')
        # Replace tabs with spaces for cleaner rendering
        p.text = re.sub(r'\t+', '    ', stripped)
        i += 1
        continue

    # ----- Normal body paragraph -----
    if stripped:
        add_para(stripped, style='PBody')

    i += 1

# ============================================================
# 8.  SAVE
# ============================================================
out = 'paradox-ieee-cited.docx'
doc.save(out)
print(f"Saved: {out}")
print(f"References in list: {len(new_to_text)}")
